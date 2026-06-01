from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = r"D:\project\clinical-rag-agent\Zhang_Chen_Individual contribution.docx"

doc = Document()

# Page margins
for section in doc.sections:
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)

# Base font
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(11)

# Title
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Individual Contribution Statement")
run.bold = True
run.font.size = Pt(18)

# Subtitle
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_run = subtitle.add_run(
    "Project: An Agentic RAG System for Medical Literature Question Answering"
)
sub_run.italic = True
sub_run.font.size = Pt(11)
sub_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()  # spacing

HEADING_COLOR = RGBColor(0x1A, 0x3D, 0x6D)


def add_heading(text: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = HEADING_COLOR


def add_body(text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.add_run(text).font.size = Pt(11)


def add_bullet(bold_lead: str, rest: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    lead = p.add_run(bold_lead)
    lead.bold = True
    lead.font.size = Pt(11)
    tail = p.add_run(rest)
    tail.font.size = Pt(11)


# ---------------- Ruoyu Zhang ----------------
add_heading("Ruoyu Zhang (ruz0002)")
add_body(
    "Ruoyu was responsible for the project end-to-end, covering topic selection, "
    "system design, full code implementation, and all written reports."
)

ruoyu_items = [
    (
        "Project topic selection and scoping: ",
        "identified the agentic medical RAG problem, defined the target benchmark "
        "(MedQA), and pivoted the project from a multimodal Med-VQA design to a "
        "text-only RAG design after finding that image-based datasets were poorly "
        "suited for retrieval evaluation.",
    ),
    (
        "System architecture and code implementation: ",
        "designed and implemented the full two-level LangGraph agent, including the "
        "top-level graph (history summarization, query rewrite and decomposition, "
        "fan-out, aggregation) and the per-sub-question agent subgraph (orchestrator, "
        "tool calling, routing, fallback, answer collection).",
    ),
    (
        "Retrieval and indexing pipeline: ",
        "hierarchical parent-child chunking algorithm; Qdrant hybrid retrieval "
        "integrating dense embeddings (all-mpnet-base-v2) and BM25 sparse search; "
        "cross-encoder reranking with ms-marco-MiniLM-L-6-v2; parent-store storage layer.",
    ),
    (
        "Agent mechanisms: ",
        "token-aware context compression with dynamic threshold; retrieval-key "
        "deduplication to prevent redundant searches; hard-limit fallback path; "
        "structured query decomposition via Pydantic schemas.",
    ),
    (
        "Prompt engineering: ",
        "authored all six domain-specific prompts (orchestrator, rewrite, summarize, "
        "compress, fallback, aggregate) with strict grounding rules and format "
        "enforcement for medical safety.",
    ),
    (
        "Evaluation and iteration: ",
        "built the automated MedQA evaluation harness; ran the v1 -> v2 -> v3 "
        "optimization cycle; analyzed failure modes and produced the qualitative "
        "case studies.",
    ),
    (
        "User interface: ",
        "Gradio chat interface with source-attribution display and streaming output.",
    ),
    (
        "Written reports: ",
        "sole author of the project proposal, mid-term report, and final report; "
        "produced all figures, tables, and LaTeX typesetting.",
    ),
]
for lead, rest in ruoyu_items:
    add_bullet(lead, rest)

doc.add_paragraph()  # spacing

# ---------------- Xu Chen ----------------
add_heading("Xu Chen (xzc0066)")
add_body("Xu was responsible for the presentation deliverables.")

xu_items = [
    (
        "Presentation slide deck: ",
        "designed and produced the slide deck summarizing the project motivation, "
        "approach, architecture, and experimental results.",
    ),
    (
        "Presentation delivery: ",
        "presented the project during class presentations.",
    ),
]
for lead, rest in xu_items:
    add_bullet(lead, rest)

doc.add_paragraph()

# Footer
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
fr = footer.add_run(
    "Both team members acknowledge and agree to the contribution breakdown above."
)
fr.italic = True
fr.font.size = Pt(10)
fr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.save(OUT)
print(f"DOCX written to: {OUT}")
